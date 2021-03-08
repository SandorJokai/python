#!/usr/bin/python3
# https://www.youtube.com/watch?v=mJEpimi_tFo&t=5363s
# apt install pyton3-pip, pip3 install python-docx, apt install python3-tk
# Before we go any further... pip install python-docx -> to this project below, we only need this lib.

from docx import Document
from docx.shared import Inches
from tkinter import *


def command():
    document = Document()
    if e1.get() != '':
        document.add_picture(e1.get(), width=Inches(1.5))
    document.add_paragraph(e2.get() + '\n' + 'Mobile: ' + e3.get() +
                           '\n' + 'Email: ' + e4.get())
    document.add_heading('ABOUT ME')
    document.add_paragraph(e5.get('1.0', END))
    document.add_heading('WORK EXPERIENCE')

    if e6.get() != '':
        document.add_paragraph().add_run(e7.get() + '-' + e8.get()).italic = True
        document.add_paragraph().add_run(e6.get()).bold = True
        document.add_paragraph(e9.get('1.0', END))
    if e10.get() != '':
        document.add_paragraph().add_run(e11.get() + '-' + e12.get()).italic = True
        document.add_paragraph().add_run(e10.get()).bold = True
        document.add_paragraph(e13.get('1.0', END))
    if e14.get() != '':
        document.add_paragraph().add_run(e15.get() + '-' + e16.get()).italic = True
        document.add_paragraph().add_run(e14.get()).bold = True
        document.add_paragraph(e17.get('1.0', END))
    document.save('CV-app.docx')

    e1.delete(0, 'end'), e2.delete(0, 'end'), e3.delete(0, 'end'), e4.delete(0, 'end')
    e5.delete('1.0', 'end'), e6.delete(0, 'end'), e9.delete('1.0', 'end')
    e10.delete(0, 'end'), e13.delete('1.0', 'end'), e14.delete(0, 'end')
    e17.delete('1.0', 'end')


root = Tk()
root.title('CV maker')
root.geometry('650x700')
photo = Label(root, text='Upload photo \n (only the path):',
              font=('Caladea', 11, 'italic'), fg='blue', bg='white').grid(row=0)
e1 = Entry(root, width=60, bd=2, bg='silver')
e1.grid(row=0, column=1)

name = Label(root, text='Full name:', fg='blue', bg='white',
             font=('Caladea', 11, 'italic')).grid(row=1)
e2 = Entry(root, width=50, bd=2, bg='silver')
e2.grid(row=1, column=1)

tel = Label(root, text='Mobile:', fg='blue', bg='white',
            font=('Caladea', 11, 'italic')).grid(row=2)
e3 = Entry(root, width=50, bd=2, bg='silver')
e3.grid(row=2, column=1)

email = Label(root, text='Email:', fg='blue', bg='white'
              , font=('Caladea', 11, 'italic')).grid(row=3)
e4 = Entry(root, width=50, bd=2, bg='silver')
e4.grid(row=3, column=1)

about_me = Label(root, text='Describe yourself:',
                 font=('Caladea', 11, 'italic'), fg='blue', bg='white').grid(row=4)
e5 = Text(root, width=50, height=4, bd=2, bg='silver')
e5.grid(row=4, column=1)

work_experience = Label(root, text='At which company?',
                        font=('Caladea', 11, 'italic'), fg='blue', bg='white').grid(row=5)
e6 = Entry(root, width=50, bd=2, bg='silver')
e6.grid(row=5, column=1)

from_date = Label(root, text='From:', font=('Caladea', 11, 'italic'),
                  fg='blue', bg='white').grid(row=6)
e7 = Spinbox(root, from_=1980, to=2020, width=8, bd=2, bg='silver')
e7.grid(row=6, column=1)

to_date = Label(root, text='To:', fg='blue', bg='white'
                , font=('Caladea', 11, 'italic')).grid(row=7)
e8 = Spinbox(root, from_=1980, to=2020, width=8, bd=2, bg='silver')
e8.grid(row=7, column=1)

tasks = Label(root, text='Main tasks and \n responsibilities?',
              font=('Caladea', 11, 'italic'), fg='blue', bg='white').grid(row=8)
e9 = Text(root, width=50, height=4, bd=2, bg='silver')
e9.grid(row=8, column=1)

work_experience2 = Label(root, text='At which company?'
                    , font=('Caladea', 11, 'italic'), fg='blue', bg='white').grid(row=9)
e10 = Entry(root, width=50, bd=2, bg='silver')
e10.grid(row=9, column=1)

from_date = Label(root, text='From:', fg='blue', bg='white'
                  , font=('Caladea', 11, 'italic')).grid(row=10)
e11 = Spinbox(root, from_=1980, to=2020, width=8, bd=2, bg='silver')
e11.grid(row=10, column=1)

to_date = Label(root, text='To:', fg='blue', bg='white'
                , font=('Caladea', 11, 'italic')).grid(row=11)
e12 = Spinbox(root, from_=1980, to=2020, width=8, bd=2, bg='silver')
e12.grid(row=11, column=1)

tasks = Label(root, text='Main tasks and \n responsibilities?', fg='blue',
              font=('Caladea', 11, 'italic'), bg='white').grid(row=12)
e13 = Text(root, width=50, height=4, bd=2, bg='silver')
e13.grid(row=12, column=1)

work_experience3 = Label(root, text='At which company?'
                , font=('Caladea', 11, 'italic'), fg='blue', bg='white').grid(row=13)
e14 = Entry(root, width=50, bd=2, bg='silver')
e14.grid(row=13, column=1)

from_date = Label(root, text='From:', fg='blue', bg='white'
                  , font=('Caladea', 11, 'italic')).grid(row=14)
e15 = Spinbox(root, from_=1980, to=2020, width=8, bd=2, bg='silver')
e15.grid(row=14, column=1)

to_date = Label(root, text='To:', fg='blue', bg='white'
                , font=('Caladea', 11, 'italic')).grid(row=15)
e16 = Spinbox(root, from_=1980, to=2020, width=8, bd=2, bg='silver')
e16.grid(row=15, column=1)

tasks = Label(root, text='Main tasks and \n responsibilities?',
              font=('Caladea', 11, 'italic'), fg='blue', bg='white').grid(row=16)
e17 = Text(root, width=50, height=4, bd=2, bg='silver')
e17.grid(row=16, column=1)

b = Button(root, text='submit', command=command, bg='white', fg='black',
           font=('Noto Sans', 14, 'italic', 'bold'), activebackground='cyan')
b.grid(row=17, column=1, ipadx=18)

root.mainloop()

'''document = Document()
print('Your cv-generator welcomes you, please answer all the questions below in order'
      ' to get as your formatted cv document. :-)')

try:
    # we start with inserting a required size picture
    document.add_picture(input('Type the full path of your passport-size picture: '),
                         width=Inches(1.5))

    name = input('what is your name: ')
    phone = input('What is your phone number: ')
    email = input('...and your email address is: ')

    document.add_paragraph(name + '\n' + 'Mobile: ' + phone +
                           '\n' + 'Email: ' + email)

    # ABOUT ME
    document.add_heading('ABOUT ME')

    me = input('Tell me a little bit about yourself...')

    document.add_paragraph(me)

    # From this section there are a few differences between this version and the previous one,
    # that is why I decided to create a modified version.

    # WORK EXPERIENCE
    document.add_heading('WORK EXPERIENCE')
    p = document.add_paragraph()
    company = ''
    from_date = ''
    to_date = ''
    tasks = ''

    anything_more = ''

    while anything_more.lower() != 'n':
        company = input('Type your previous company(ies) you work with: ')
        from_date = input('From what year you started to work there? ')
        to_date = input('When did you leave that company? ')
        tasks = input('What were your main tasks and responsibilities? ')
        p.add_run(from_date + '-' + to_date + '\n').italic = True
        p.add_run(company + '\n').bold = True
        p.add_run(tasks + '\n' + '\n')
        anything_more = input('Is there anything more experience you\'d like '
                              'to share? (y)es or (n)o? ')

    document.save('emoke-CV.docx')
    print('Your nice and simple CV has just been created. It\'s time to apply for a job.'
          ' Good luck with that. :)')
except BaseException:
    print('Something wrong, please try again.')
'''
