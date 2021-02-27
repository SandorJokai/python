# Before we go any further... pip install python-docx -> to this project below, we only need this lib.

from docx import Document
from docx.shared import Inches

document = Document()
print('Your cv-generator welcomes you, please answer all the questions below in order'
      ' to get as your formatted cv document. :-)')

try:
    # we start with inserting a required size picture
    document.add_picture(input('Type the full path of your passport-size picture: '),
                         width=Inches(1.5))

    name = input('what is your name: ')
    phone = input('What is your phone number: ')
    email = input('...and your email address is: ')

    document.add_paragraph(name + '\n' + 'Mobile: ' + phone +
                           '\n' + 'Email: ' + email)

    # ABOUT ME
    document.add_heading('ABOUT ME')

    me = input('Tell me a little bit about yourself...')

    document.add_paragraph(me)

    # From this section there are a few differences between this version and the previous one,
    # that is why I decided to create a modified version.

    # WORK EXPERIENCE
    document.add_heading('WORK EXPERIENCE')
    p = document.add_paragraph()
    company = ''
    from_date = ''
    to_date = ''
    tasks = ''

    anything_more = ''

    while anything_more.lower() != 'n':
        company = input('Type your previous company(ies) you work with: ')
        from_date = input('From what year you started to work there? ')
        to_date = input('When did you leave that company? ')
        tasks = input('What were your main tasks and responsibilities? ')
        p.add_run(from_date + '-' + to_date + '\n').italic = True
        p.add_run(company + '\n').bold = True
        p.add_run(tasks + '\n' + '\n')
        anything_more = input('Is there anything more experience you\'d like '
                              'to share? (y)es or (n)o? ')

    document.save('emoke-CV.docx')
    print('Your nice and simple CV has just been created. It\'s time to apply for a job.'
          ' Good luck with that. :)')
except BaseException:
    print('Something wrong, please try again.')
