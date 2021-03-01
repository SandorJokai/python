#/usr/bin/python3
from docx import Document
from docx.shared import Inches
from tkinter import *


def command():
    document = Document()
    document.add_picture(e1.get(), width=Inches(1.5))
    document.add_paragraph(e2.get() + '\n' + 'Mobile: ' + e3.get() +
                           '\n' + 'Email: ' + e4.get())
    document.add_heading('ABOUT ME')
    document.add_paragraph(e5.get('1.0', END))
    document.add_heading('WORK EXPERIENCE')

    if e6.get() != '':
        document.add_paragraph().add_run(e7.get() + '-' + e8.get()).italic = True
        document.add_paragraph().add_run(e6.get()).bold = True
        document.add_paragraph(e9.get('1.0', END))
    if e10.get() != '':
        document.add_paragraph().add_run(e11.get() + '-' + e12.get()).italic = True
        document.add_paragraph().add_run(e10.get()).bold = True
        document.add_paragraph(e13.get('1.0', END))
    if e14.get() != '':
        document.add_paragraph().add_run(e15.get() + '-' + e16.get()).italic = True
        document.add_paragraph().add_run(e14.get()).bold = True
        document.add_paragraph(e17.get('1.0', END))
    document.save('CV-app.docx')


root = Tk()
root.title('CV maker')
root.geometry('650x700')
photo = Label(root, text='Upload photo \n (only the path):').grid(row=0)
e1 = Entry(root, width=60)
e1.grid(row=0, column=1)
name = Label(root, text='Full name:').grid(row=1)
e2 = Entry(root, width=40)
e2.grid(row=1, column=1)
tel = Label(root, text='Mobile:').grid(row=2)
e3 = Entry(root, width=40)
e3.grid(row=2, column=1)
email = Label(root, text='Email:').grid(row=3)
e4 = Entry(root, width=40)
e4.grid(row=3, column=1)
about_me = Label(root, text='Describe yourself:').grid(row=4)
e5 = Text(root, width=40, height=5)
e5.grid(row=4, column=1)

work_experience = Label(root, text='At which company?').grid(row=5)
e6 = Entry(root, width=40)
e6.grid(row=5, column=1)
from_date = Label(root, text='From:   ').grid(row=6)
e7 = Spinbox(root, from_=1980, to=2020, width=8)
e7.grid(row=6, column=1)
to_date = Label(root, text='To:   ').grid(row=7)
e8 = Spinbox(root, from_=1980, to=2020, width=8)
e8.grid(row=7, column=1)
tasks = Label(root, text='Main tasks and \n responsibilities?   ').grid(row=8)
e9 = Text(root, width=40, height=5)
e9.grid(row=8, column=1)

work_experience2 = Label(root, text='At which company?').grid(row=9)
e10 = Entry(root, width=40)
e10.grid(row=9, column=1)
from_date = Label(root, text='From:   ').grid(row=10)
e11 = Spinbox(root, from_=1980, to=2020, width=8)
e11.grid(row=10, column=1)
to_date = Label(root, text='To:   ').grid(row=11)
e12 = Spinbox(root, from_=1980, to=2020, width=8)
e12.grid(row=11, column=1)
tasks = Label(root, text='Main tasks and \n responsibilities?   ').grid(row=12)
e13 = Text(root, width=40, height=5)
e13.grid(row=12, column=1)

work_experience3 = Label(root, text='At which company?').grid(row=13)
e14 = Entry(root, width=40)
e14.grid(row=13, column=1)
from_date = Label(root, text='From:   ').grid(row=14)
e15 = Spinbox(root, from_=1980, to=2020, width=8)
e15.grid(row=14, column=1)
to_date = Label(root, text='To:   ').grid(row=15)
e16 = Spinbox(root, from_=1980, to=2020, width=8)
e16.grid(row=15, column=1)
tasks = Label(root, text='Main tasks and \n responsibilities?   ').grid(row=16)
e17 = Text(root, width=40, height=5)
e17.grid(row=16, column=1)

b = Button(root, text='commit', command=command)
b.grid(row=17, column=1)

root.mainloop()
